"""
Document Processor - Handles reading and chunking different document types.
Supports: PDF, DOCX, XLSX, TXT
"""
import os
import re
import math
import hashlib
from collections import Counter
from typing import List, Dict
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from openpyxl import load_workbook
from backend.config import (
    CHUNK_SIZE,
    CHUNK_OVERLAP,
    MIN_CHUNK_CHARS,
    ENABLE_CHUNK_DEDUP,
    MAX_CHUNKS_PER_SECTION,
    PDF_MARGIN_REPEAT_RATIO,
    SUPPORTED_EXTENSIONS,
)


_SENTENCE_SPLIT_RE = re.compile(r"(?<=[\.!?;:។])\s+")
_BULLET_RE = re.compile(r"^(?:[-*•]|\d+[\)\.]|[a-zA-Z][\)\.])\s+")
_TABLE_SEPARATOR_RE = re.compile(r"\s*\|\s*")
_PAGE_NUMBER_RE = re.compile(r"^(?:page|trang)?\s*\d+(?:\s*/\s*\d+)?$", re.IGNORECASE)


def _normalize_line(line: str) -> str:
    """Normalize a single line for comparison and chunk assembly."""
    return re.sub(r"\s+", " ", line).strip()


def _line_signature(line: str) -> str:
    """Build a comparison signature for duplicate header/footer detection."""
    return _normalize_line(line).casefold()


def _is_page_number_line(line: str) -> bool:
    """Detect standalone page-number footer/header lines."""
    clean = _normalize_line(line)
    if not clean:
        return False
    return bool(_PAGE_NUMBER_RE.match(clean) or re.fullmatch(r"\d+", clean))


def _split_lines(text: str) -> List[str]:
    """Split text into cleaned, non-empty lines."""
    return [_normalize_line(line) for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n") if _normalize_line(line)]


def _can_be_margin_line(line: str) -> bool:
    """Heuristic filter to avoid removing real content as page header/footer."""
    clean = _normalize_line(line)
    if not clean:
        return False
    if _is_page_number_line(clean):
        return True
    if len(clean) > 90:
        return False
    if _TABLE_SEPARATOR_RE.search(clean) or _BULLET_RE.match(clean):
        return False
    if clean.endswith((".", "!", "?", ";")) and len(clean.split()) > 6:
        return False
    return True


def _collect_repeated_margin_signatures(page_lines: List[List[str]]) -> set:
    """Detect repeated header/footer lines across pages conservatively."""
    if len(page_lines) < 2:
        return set()

    top_counts = Counter()
    bottom_counts = Counter()
    sample_line_by_signature: Dict[str, str] = {}

    for lines in page_lines:
        if not lines:
            continue
        top_candidates = lines[:2]
        bottom_candidates = lines[-2:] if len(lines) > 2 else []

        top_seen = set()
        for line in top_candidates:
            if not _can_be_margin_line(line):
                continue
            signature = _line_signature(line)
            if not signature or signature in top_seen:
                continue
            top_counts[signature] += 1
            top_seen.add(signature)
            sample_line_by_signature.setdefault(signature, line)

        bottom_seen = set()
        for line in bottom_candidates:
            if not _can_be_margin_line(line):
                continue
            signature = _line_signature(line)
            if not signature or signature in bottom_seen:
                continue
            bottom_counts[signature] += 1
            bottom_seen.add(signature)
            sample_line_by_signature.setdefault(signature, line)

    threshold = max(2, math.ceil(len(page_lines) * PDF_MARGIN_REPEAT_RATIO))
    repeated_signatures = set()
    for signature, count in top_counts.items():
        line = sample_line_by_signature.get(signature, "")
        if count >= threshold and len(line.split()) <= 14:
            repeated_signatures.add(signature)
    for signature, count in bottom_counts.items():
        line = sample_line_by_signature.get(signature, "")
        if count >= threshold and len(line.split()) <= 14:
            repeated_signatures.add(signature)

    return repeated_signatures


def _remove_repeated_margin_lines(lines: List[str], repeated_signatures: set) -> List[str]:
    """Remove page headers/footers that repeat across pages."""
    if not repeated_signatures or not lines:
        return lines

    cleaned = []
    last_index = len(lines) - 1
    for index, line in enumerate(lines):
        signature = _line_signature(line)
        near_top = index < 2
        near_bottom = index >= max(0, last_index - 1)
        if signature in repeated_signatures and (near_top or near_bottom or _is_page_number_line(line)):
            continue
        cleaned.append(line)

    if not cleaned:
        # Safety fallback: never drop a full page.
        return [line for line in lines if not _is_page_number_line(line)] or lines
    return cleaned


def extract_text_from_pdf(file_path: str) -> List[Dict]:
    """Extract text from PDF file, split page content into semantic sections."""
    reader = PdfReader(file_path)
    raw_pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            raw_pages.append({
                "page": i + 1,
                "text": text.strip(),
            })

    if not raw_pages:
        return []

    page_lines = [_split_lines(page["text"]) for page in raw_pages]
    repeated_signatures = _collect_repeated_margin_signatures(page_lines)

    pages = []
    for page, lines in zip(raw_pages, page_lines):
        cleaned_lines = _remove_repeated_margin_lines(lines, repeated_signatures)
        cleaned_text = "\n".join(cleaned_lines).strip()
        if not cleaned_text:
            continue

        sections = _split_into_blocks(cleaned_text)
        if not sections:
            sections = [cleaned_text]

        for section_index, section_text in enumerate(sections):
            pages.append({
                "text": section_text.strip(),
                "metadata": {
                    "page": page["page"],
                    "page_block_index": section_index,
                }
            })
    return pages


def extract_text_from_docx(file_path: str) -> List[Dict]:
    """Extract text from Word document."""
    doc = DocxDocument(file_path)
    full_text = []
    for para in doc.paragraphs:
        paragraph_text = para.text.strip()
        if not paragraph_text:
            continue

        style_name = getattr(getattr(para, "style", None), "name", "") or ""
        style_name_lower = style_name.lower()
        if style_name_lower.startswith("heading"):
            heading_level_match = re.search(r"(\d+)", style_name_lower)
            heading_level = int(heading_level_match.group(1)) if heading_level_match else 1
            heading_prefix = "#" * max(1, min(6, heading_level))
            full_text.append(f"{heading_prefix} {paragraph_text}")
        else:
            full_text.append(paragraph_text)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                full_text.append(row_text)

    text = "\n".join(full_text)
    if text:
        return [{"text": text, "metadata": {}}]
    return []


def extract_text_from_xlsx(file_path: str) -> List[Dict]:
    """Extract text from Excel file, sheet by sheet."""
    sheets = []
    wb = load_workbook(file_path, read_only=True, data_only=True)
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row_index, row in enumerate(ws.iter_rows(values_only=True), start=1):
            row_text = " | ".join(str(cell) for cell in row if cell is not None)
            if row_text.strip():
                rows.append(f"Row {row_index}: {row_text}")
        if rows:
            text = "\n".join(rows)
            sheets.append({
                "text": text,
                "metadata": {"sheet": sheet_name}
            })
    wb.close()
    return sheets


def extract_text_from_txt(file_path: str) -> List[Dict]:
    """Extract text from plain text file."""
    encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]
    for enc in encodings:
        try:
            with open(file_path, "r", encoding=enc) as f:
                text = f.read().strip()
            if text:
                return [{"text": text, "metadata": {}}]
            return []
        except (UnicodeDecodeError, UnicodeError):
            continue
    return []


EXTRACTORS = {
    ".pdf": extract_text_from_pdf,
    ".docx": extract_text_from_docx,
    ".xlsx": extract_text_from_xlsx,
    ".txt": extract_text_from_txt,
}


def _normalize_text(text: str) -> str:
    """Normalize whitespace while preserving paragraph boundaries and list structure."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"(?<=\w)-\n(?=\w)", "", text)
    lines = [_normalize_line(line) for line in text.split("\n")]
    blocks = []
    current = []

    for line in lines:
        if not line:
            if current:
                blocks.append(current)
                current = []
            continue
        current.append(line)

    if current:
        blocks.append(current)

    formatted_blocks = []
    for block_lines in blocks:
        if len(block_lines) == 1:
            formatted_blocks.append(block_lines[0])
            continue

        if all(_is_table_like_line(line) for line in block_lines):
            formatted_blocks.append("\n".join(block_lines))
            continue

        formatted_blocks.append(" ".join(block_lines))

    return "\n\n".join(formatted_blocks)


def _is_table_like_line(line: str) -> bool:
    """Return True when a line looks like a table row or bullet entry."""
    return bool(_BULLET_RE.match(line) or _TABLE_SEPARATOR_RE.search(line))


def _is_heading_block(block: str) -> bool:
    """Heuristically detect headings so they can be kept with following content."""
    clean = block.strip()
    if not clean:
        return False
    if len(clean) > 120:
        return False
    if clean.startswith("#"):
        return True
    if clean.endswith(":"):
        return True

    words = clean.split()
    if len(words) <= 12 and clean.upper() == clean and any(ch.isalpha() for ch in clean):
        return True

    if len(words) <= 10:
        capitalized = sum(1 for word in words if word and word[0].isupper())
        if capitalized >= max(2, len(words) // 2) and not clean.endswith("."):
            return True

    return bool(re.match(r"^(?:\d+(?:\.\d+)*|[IVXLC]+)[\)\.]\s+", clean, re.IGNORECASE))


def _split_into_blocks(text: str) -> List[str]:
    """Split normalized text into semantic blocks while preserving lists and tables."""
    normalized = _normalize_text(text)
    if not normalized:
        return []

    raw_blocks = [block.strip() for block in re.split(r"\n\n+", normalized) if block.strip()]
    merged_blocks: List[str] = []
    pending_headings: List[str] = []

    for block in raw_blocks:
        if _is_heading_block(block):
            pending_headings.append(block)
            continue

        if pending_headings:
            merged_blocks.append("\n".join(pending_headings + [block]))
            pending_headings = []
        else:
            merged_blocks.append(block)

    if pending_headings:
        merged_blocks.extend(pending_headings)

    return merged_blocks


def _normalize_chunk_signature(text: str) -> str:
    """Normalize chunk text for deduplication comparison."""
    lowered = text.casefold()
    lowered = re.sub(r"\s+", " ", lowered)
    lowered = re.sub(r"[^\w\s]", "", lowered)
    return lowered.strip()


def _is_low_value_chunk(text: str) -> bool:
    """Detect chunks that carry almost no retrievable signal."""
    clean = text.strip()
    if not clean:
        return True
    if len(clean) < 20 and len(re.findall(r"[A-Za-zÀ-ỹ0-9]", clean)) < 8:
        return True
    return False


def _merge_small_chunks(chunks: List[str], chunk_size: int, min_chunk_chars: int) -> List[str]:
    """Merge very short chunks into neighbors to improve retrieval quality."""
    if not chunks:
        return []

    merged: List[str] = []
    for chunk in chunks:
        chunk = chunk.strip()
        if not chunk:
            continue

        if not merged:
            merged.append(chunk)
            continue

        if len(chunk) < min_chunk_chars:
            candidate = f"{merged[-1]}\n\n{chunk}".strip()
            if len(candidate) <= int(chunk_size * 1.25):
                merged[-1] = candidate
                continue

        merged.append(chunk)

    return merged


def _post_process_chunks(raw_chunks: List[str], chunk_size: int, min_chunk_chars: int, enable_dedup: bool) -> List[str]:
    """Clean, merge, and deduplicate chunks before indexing."""
    cleaned = [chunk.strip() for chunk in raw_chunks if chunk and chunk.strip()]
    cleaned = [chunk for chunk in cleaned if not _is_low_value_chunk(chunk)]
    cleaned = _merge_small_chunks(cleaned, chunk_size, min_chunk_chars)

    if not enable_dedup:
        return cleaned

    deduped: List[str] = []
    seen_signatures = set()
    for chunk in cleaned:
        signature = _normalize_chunk_signature(chunk)
        if not signature or signature in seen_signatures:
            continue
        seen_signatures.add(signature)
        deduped.append(chunk)
    return deduped


def _infer_content_kind(text: str) -> str:
    """Tag chunk type to support debugging and future reranking policies."""
    if "|" in text:
        return "table"
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if lines and sum(1 for line in lines if _BULLET_RE.match(line)) >= max(1, len(lines) // 2):
        return "list"
    if lines and _is_heading_block(lines[0]) and len(lines) <= 3:
        return "heading"
    return "paragraph"


def _split_long_unit(unit: str, max_size: int) -> List[str]:
    """Split oversized text unit by sentence boundaries, fallback to hard split."""
    if len(unit) <= max_size:
        return [unit]

    if "\n" in unit:
        lines = [line.strip() for line in unit.splitlines() if line.strip()]
        if len(lines) > 1:
            split_lines = []
            current_line = ""
            for line in lines:
                candidate = f"{current_line}\n{line}".strip() if current_line else line
                if len(candidate) <= max_size:
                    current_line = candidate
                    continue
                if current_line:
                    split_lines.append(current_line)
                current_line = line
            if current_line:
                split_lines.append(current_line)
            if split_lines:
                return split_lines

    sentence_parts = [part.strip() for part in _SENTENCE_SPLIT_RE.split(unit) if part.strip()]
    if len(sentence_parts) <= 1:
        return [unit[i:i + max_size].strip() for i in range(0, len(unit), max_size) if unit[i:i + max_size].strip()]

    result = []
    current = ""
    for sentence in sentence_parts:
        candidate = f"{current} {sentence}".strip() if current else sentence
        if len(candidate) <= max_size:
            current = candidate
            continue

        if current:
            result.append(current)
        if len(sentence) <= max_size:
            current = sentence
        else:
            hard_parts = [sentence[i:i + max_size].strip() for i in range(0, len(sentence), max_size) if sentence[i:i + max_size].strip()]
            result.extend(hard_parts[:-1])
            current = hard_parts[-1] if hard_parts else ""

    if current:
        result.append(current)

    return result


def _collect_overlap_blocks(blocks: List[str], overlap: int) -> List[str]:
    """Collect the smallest suffix of blocks that still satisfies the overlap budget."""
    if overlap <= 0 or not blocks:
        return []

    overlap_blocks: List[str] = []
    overlap_len = 0

    for block in reversed(blocks):
        block_len = len(block)
        added_len = block_len if not overlap_blocks else block_len + 2
        if overlap_blocks and overlap_len + added_len > overlap:
            break
        overlap_blocks.append(block)
        overlap_len += added_len
        if overlap_len >= overlap:
            break

    return list(reversed(overlap_blocks))


def _build_chunks_from_units(units: List[str], chunk_size: int, overlap: int) -> List[str]:
    """Assemble units into semantic, overlap-aware chunks."""
    if not units:
        return []

    chunks = []
    current_units: List[str] = []

    for unit in units:
        if len(unit) > chunk_size:
            split_units = _split_long_unit(unit, chunk_size)
        else:
            split_units = [unit]

        for split_unit in split_units:
            split_unit = split_unit.strip()
            if not split_unit:
                continue

            candidate_units = current_units + [split_unit]
            candidate_text = "\n\n".join(candidate_units).strip()
            if current_units and len(candidate_text) > chunk_size:
                chunks.append("\n\n".join(current_units).strip())
                overlap_units = _collect_overlap_blocks(current_units, overlap)
                current_units = overlap_units
                candidate_units = current_units + [split_unit]
                candidate_text = "\n\n".join(candidate_units).strip()

                if candidate_text and len(candidate_text) > chunk_size:
                    if overlap_units:
                        candidate_units = [split_unit]
                        candidate_text = split_unit
                    else:
                        # Hard split on the current unit if it still does not fit after semantic splitting.
                        chunks.extend(_split_long_unit(split_unit, chunk_size))
                        current_units = []
                        continue

            if candidate_text and len(candidate_text) > chunk_size and not current_units:
                # Hard split on the current unit if it still does not fit after semantic splitting.
                chunks.extend(_split_long_unit(split_unit, chunk_size))
                current_units = []
                continue

            current_units = candidate_units

    if current_units:
        chunks.append("\n\n".join(current_units).strip())

    return chunks


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """Split text into semantically coherent overlapping chunks."""
    blocks = _split_into_blocks(text)
    if not blocks:
        return []

    expanded_units: List[str] = []
    for unit in blocks:
        expanded_units.extend(_split_long_unit(unit, chunk_size))

    raw_chunks = _build_chunks_from_units(expanded_units, chunk_size, overlap)
    return _post_process_chunks(raw_chunks, chunk_size, MIN_CHUNK_CHARS, ENABLE_CHUNK_DEDUP)


def process_document(file_path: str) -> List[Dict]:
    """
    Process a document file and return a list of chunks with metadata.
    
    Returns:
        List of dicts: [{"text": "...", "metadata": {"source": "file.pdf", "page": 1, ...}}, ...]
    """
    ext = os.path.splitext(file_path)[1].lower()
    filename = os.path.basename(file_path)
    # Strip doc_id prefix (format: "abc12345_originalname.ext")
    if "_" in filename:
        filename = filename.split("_", 1)[1]

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {SUPPORTED_EXTENSIONS}")

    extractor = EXTRACTORS.get(ext)
    if not extractor:
        raise ValueError(f"No extractor for file type: {ext}")

    # Extract text sections
    sections = extractor(file_path)

    # Chunk each section
    all_chunks = []
    global_chunk_index = 0
    for section_index, section in enumerate(sections):
        text = section.get("text", "")
        if not text or not text.strip():
            continue

        meta = {**section.get("metadata", {})}
        meta["source"] = filename
        meta["section_index"] = section_index

        chunks = chunk_text(text)
        if MAX_CHUNKS_PER_SECTION > 0:
            chunks = chunks[:MAX_CHUNKS_PER_SECTION]

        for i, chunk in enumerate(chunks):
            chunk_hash = hashlib.sha1(chunk.encode("utf-8", errors="ignore")).hexdigest()[:16]
            chunk_meta = {
                **meta,
                "section_chunk_index": i,
                "chunk_index": global_chunk_index,
                "chunk_char_count": len(chunk),
                "content_kind": _infer_content_kind(chunk),
                "chunk_hash": chunk_hash,
            }
            all_chunks.append({
                "text": chunk,
                "metadata": chunk_meta
            })
            global_chunk_index += 1

    return all_chunks
